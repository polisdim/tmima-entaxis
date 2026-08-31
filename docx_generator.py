#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Word (.docx) Document Generator for Inclusion Class (Τμήμα Ένταξης)
Populates official school templates based on live student data, Bloom levels,
interventions, 4-level IEP rubrics, and Duval/Brousseau didactic frameworks.
School: ΔΗΜ.Ω.Σ. Γυμνάσιο Ξάνθης
Teacher: Δημήτριος Πολυχρόνης (ΠΕ03.ΕΑΕ)
"""

import os
import copy
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from datetime import datetime

TEMPLATES_BASE_DIR = r"c:\Users\polis\OneDrive - Democritus University of Thrace\Υλικά\ΕΑΕ\Ένταξη"

RUBRIC_LEVEL_MAP = {
    1: "1. Αρχικό (Πλήρης καθοδήγηση)",
    2: "2. Αναδυόμενο (Μερική υποστήριξη)",
    3: "3. Ικανό (Αυτόνομη επίλυση)",
    4: "4. Γενικευμένο (Πλήρης κατάκτηση)"
}

def get_template_file(template_filename, gender="Κορίτσι"):
    """Finds the template file in gender-specific folder with fallback. Returns None if not found."""
    if not os.path.exists(TEMPLATES_BASE_DIR):
        return None

    folder_name = "0. Αγόρι" if (gender and "Αγόρι" in gender) else "0. Κορίτσι"
    path = os.path.join(TEMPLATES_BASE_DIR, folder_name, template_filename)
    if os.path.exists(path):
        return path
    
    fallback_path = os.path.join(TEMPLATES_BASE_DIR, "0. Αγόρι", template_filename)
    if os.path.exists(fallback_path):
        return fallback_path
        
    fallback_path2 = os.path.join(TEMPLATES_BASE_DIR, "0. Κορίτσι", template_filename)
    if os.path.exists(fallback_path2):
        return fallback_path2

    return None

def create_standalone_docx(student, observations, doc_type="iep"):
    """Creates a beautifully styled Word document formatted for Greek Inclusion Class."""
    doc = docx.Document()
    full_name = f"{student.get('name', '')} {student.get('surname', '')}".strip()
    
    # Title
    title = doc.add_heading("ΤΜΗΜΑ ΕΝΤΑΞΗΣ - ΔΗΜ.Ω.Σ. ΓΥΜΝΑΣΙΟ ΞΑΝΘΗΣ", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc_titles = {
        "iep": "ΕΞΑΤΟΜΙΚΕΥΜΕΝΟ ΠΡΟΓΡΑΜΜΑ ΕΚΠΑΙΔΕΥΣΗΣ (Ε.Π.Ε.)",
        "aa": "ΑΡΧΙΚΗ ΔΙΑΓΝΩΣΤΙΚΗ ΑΞΙΟΛΟΓΗΣΗ",
        "rubrics": "ΕΝΔΙΑΜΕΣΗ ΑΞΙΟΛΟΓΗΣΗ & ΡΟΥΜΠΡΙΚΕΣ ΔΙΑΒΑΘΜΙΣΗΣ",
        "ta": "ΤΕΛΙΚΗ ΕΚΘΕΣΗ ΑΠΟΤΙΜΗΣΗΣ & ΜΕΤΑΒΑΣΗΣ"
    }
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run(f"{doc_titles.get(doc_type, 'Ε.Π.Ε.')} - ΣΧΟΛΙΚΟ ΕΤΟΣ 2026-2027\nΕκπαιδευτικός Ε.Α.Ε.: Δημήτριος Πολυχρόνης (ΠΕ03.ΕΑΕ)")
    r.bold = True
    
    # 1. Demographic Info
    doc.add_heading("1. Δημογραφικά Στοιχεία & Διαγνωστικό Ιστορικό", level=1)
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    father = student.get('parent_father', {})
    mother = student.get('parent_mother', {})
    diag = student.get('diagnosis_info', {})
    
    table.rows[0].cells[0].text = "Ονοματεπώνυμο Μαθητή/τριας:"
    table.rows[0].cells[1].text = full_name
    table.rows[1].cells[0].text = "Τάξη / Τμήμα / Α.Μ.:"
    table.rows[1].cells[1].text = f"{student.get('grade', '')} {student.get('class_section', '')} (Α.Μ: {student.get('registration_no', '-')})"
    table.rows[2].cells[0].text = "Διάγνωση & Φορέας:"
    table.rows[2].cells[1].text = f"{diag.get('diagnosis_type', student.get('diagnosis', 'Ειδική Μαθησιακή Δυσκολία'))} ({diag.get('authority', 'ΚΕΔΑΣΥ Ξάνθης')})"
    table.rows[3].cells[0].text = "Στοιχεία Πατέρα:"
    table.rows[3].cells[1].text = f"{father.get('name', '-')} (Τηλ: {father.get('phone', '-')})"
    table.rows[4].cells[0].text = "Στοιχεία Μητέρας:"
    table.rows[4].cells[1].text = f"{mother.get('name', '-')} (Τηλ: {mother.get('phone', '-')})"
    
    # 2. Math Profile
    doc.add_heading("2. Μαθησιακό & Γνωστικό Προφίλ στα Μαθηματικά", level=1)
    mprof = student.get('math_profile', {})
    doc.add_paragraph(f"• Επίπεδο Μαθηματικού Άγχους: {mprof.get('math_anxiety', 'Μέτριο')}")
    doc.add_paragraph(f"• Προτιμώμενο Μαθησιακό Στυλ: {mprof.get('learning_style', 'Οπτικό / Πολυαισθητηριακό')}")
    if mprof.get('strengths'):
        doc.add_paragraph(f"• Δυνατά Σημεία: {mprof.get('strengths')}")
    if mprof.get('weaknesses'):
        doc.add_paragraph(f"• Κύρια Εμπόδια / Τομείς Δυσκολίας: {mprof.get('weaknesses')}")
    if mprof.get('effective_strategies'):
        doc.add_paragraph(f"• Αποτελεσματικές Διδακτικές Στρατηγικές: {mprof.get('effective_strategies')}")
        
    # 3. IEP Goals & 4-Level Rubric
    doc.add_heading("3. Στοχοθεσία Ε.Π.Ε. & Διαβαθμισμένη Ρουμπρίκα 4 Επιπέδων", level=1)
    iep_targets = student.get('iep_targets', [])
    if iep_targets:
        t_table = doc.add_table(rows=1 + len(iep_targets), cols=4)
        t_table.style = 'Table Grid'
        t_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        t_table.rows[0].cells[0].text = "Τομέας"
        t_table.rows[0].cells[1].text = "Διδακτικός Στόχος Ε.Π.Ε."
        t_table.rows[0].cells[2].text = "Στάθμη Ρουμπρίκας (1-4)"
        t_table.rows[0].cells[3].text = "Κατάσταση"
        for i, t in enumerate(iep_targets, 1):
            t_table.rows[i].cells[0].text = t.get('area', 'Μαθηματικά')
            t_table.rows[i].cells[1].text = t.get('target', '-')
            r_lvl = t.get('rubric_level', 2)
            t_table.rows[i].cells[2].text = RUBRIC_LEVEL_MAP.get(r_lvl, f"Επίπεδο {r_lvl}")
            t_table.rows[i].cells[3].text = t.get('status', 'Σε εξέλιξη')
            
    # 4. History of Interventions
    doc.add_heading("4. Ιστορικό Διδακτικών Παρεμβάσεων στο Τμήμα Ένταξης", level=1)
    st_id = student.get('id')
    st_obs = [o for o in observations if o.get('student_id') == st_id or o.get('student_name') == student.get('name')]
    if st_obs:
        obs_table = doc.add_table(rows=1 + len(st_obs), cols=5)
        obs_table.style = 'Table Grid'
        obs_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        obs_table.rows[0].cells[0].text = "Ημερομηνία"
        obs_table.rows[0].cells[1].text = "Τομέας & Bloom"
        obs_table.rows[0].cells[2].text = "Αναπαράσταση (Duval)"
        obs_table.rows[0].cells[3].text = "Στρατηγική & Αποτέλεσμα"
        obs_table.rows[0].cells[4].text = "Περιγραφή Παρατήρησης"
        for i, o in enumerate(st_obs, 1):
            obs_table.rows[i].cells[0].text = (o.get('timestamp') or '')[:10]
            obs_table.rows[i].cells[1].text = f"{o.get('domain_name','')}\n({o.get('bloom_name','')})"
            obs_table.rows[i].cells[2].text = o.get('duval_name', o.get('duval_register', '-'))
            obs_table.rows[i].cells[3].text = f"{o.get('strategy_name','')}\n[{o.get('outcome_code','+')}]"
            obs_table.rows[i].cells[4].text = o.get('raw_text', '')
            
    return doc

def generate_iep_docx(student, observations, output_dir):
    """Generates the official IEP (.docx) document."""
    gender = student.get('gender', 'Κορίτσι')
    tmpl_path = get_template_file("3. ΕΠΩΝΥΜΟ ΟΝΟΜΑ ΕΠΕ.docx", gender)
    if not tmpl_path:
        doc = create_standalone_docx(student, observations, "iep")
    else:
        doc = docx.Document(tmpl_path)
    
    if len(doc.tables) > 0:
        t0 = doc.tables[0]
        parts = student.get('name', '').split(' ', 1)
        first_name = parts[0]
        last_name = parts[1] if len(parts) > 1 else ''
        
        if len(t0.rows) > 1 and len(t0.rows[1].cells) > 3:
            t0.rows[1].cells[1].text = first_name
            t0.rows[1].cells[3].text = last_name or student.get('name', '')
        
        if len(t0.rows) > 2 and len(t0.rows[2].cells) > 3:
            t0.rows[2].cells[1].text = f"{student.get('grade', '')} {student.get('class_section', '')}".strip()
            t0.rows[2].cells[3].text = "2026-2027"
        
        if len(t0.rows) > 3 and len(t0.rows[3].cells) > 3:
            t0.rows[3].cells[1].text = student.get('diagnosis', 'Ειδική Μαθησιακή Δυσκολία')
            t0.rows[3].cells[3].text = "2026-2027"

    os.makedirs(output_dir, exist_ok=True)
    st_clean_name = student.get('name', 'Μαθητής').replace(' ', '_')
    out_filename = f"2_ΕΠΕ_{st_clean_name}_2026-2027.docx"
    out_path = os.path.join(output_dir, out_filename)
    doc.save(out_path)
    return out_path

def generate_initial_assessment_docx(student, observations, output_dir):
    """Generates the Initial Assessment (.docx) document."""
    gender = student.get('gender', 'Κορίτσι')
    tmpl_path = get_template_file("2. ΕΠΩΝΥΜΟ ΟΝΟΜΑ ΑΑ.docx", gender)
    if not tmpl_path:
        doc = create_standalone_docx(student, observations, "aa")
    else:
        doc = docx.Document(tmpl_path)
    
    if len(doc.tables) > 0:
        t0 = doc.tables[0]
        parts = student.get('name', '').split(' ', 1)
        first_name = parts[0]
        last_name = parts[1] if len(parts) > 1 else ''
        
        if len(t0.rows) > 1 and len(t0.rows[1].cells) > 5:
            t0.rows[1].cells[1].text = first_name
            t0.rows[1].cells[5].text = last_name or student.get('name', '')
        if len(t0.rows) > 2 and len(t0.rows[2].cells) > 3:
            t0.rows[2].cells[1].text = student.get('grade', "Β' Γυμνασίου")
            t0.rows[2].cells[3].text = student.get('class_section', "Β1")
        if len(t0.rows) > 3 and len(t0.rows[3].cells) > 6:
            t0.rows[3].cells[1].text = student.get('diagnosis', 'Ειδική Μαθησιακή Δυσκολία')
            t0.rows[3].cells[6].text = "2026-2027"

    os.makedirs(output_dir, exist_ok=True)
    st_clean_name = student.get('name', 'Μαθητής').replace(' ', '_')
    out_filename = f"1_Αρχική_Αξιολόγηση_{st_clean_name}.docx"
    out_path = os.path.join(output_dir, out_filename)
    doc.save(out_path)
    return out_path

def generate_rubrics_docx(student, observations, output_dir):
    """Generates the Mid-Year Evaluation / Rubrics (.docx) document."""
    gender = student.get('gender', 'Κορίτσι')
    tmpl_path = get_template_file("4. ΕΠΩΝΥΜΟ ΟΝΟΜΑ ΑΞΙΟΛΟΓΗΣΗ.docx", gender)
    if not tmpl_path:
        doc = create_standalone_docx(student, observations, "rubrics")
    else:
        doc = docx.Document(tmpl_path)
    
    for t in doc.tables:
        if len(t.rows) > 1:
            parts = student.get('name', '').split(' ', 1)
            if len(t.rows[1].cells) > 1:
                t.rows[1].cells[1].text = parts[0]
            if len(t.rows[1].cells) > 4:
                t.rows[1].cells[4].text = parts[1] if len(parts) > 1 else ''

    os.makedirs(output_dir, exist_ok=True)
    st_clean_name = student.get('name', 'Μαθητής').replace(' ', '_')
    out_filename = f"3_Ενδιάμεση_Αξιολόγηση_{st_clean_name}.docx"
    out_path = os.path.join(output_dir, out_filename)
    doc.save(out_path)
    return out_path

def generate_final_evaluation_docx(student, observations, output_dir):
    """Generates the Final Summative Evaluation Report (.docx) document."""
    gender = student.get('gender', 'Κορίτσι')
    tmpl_path = get_template_file("6. ΕΠΩΝΥΜΟ ΟΝΟΜΑ ΤΑ.docx", gender)
    if not tmpl_path:
        doc = create_standalone_docx(student, observations, "ta")
    else:
        doc = docx.Document(tmpl_path)
    
    if len(doc.tables) > 0:
        t0 = doc.tables[0]
        parts = student.get('name', '').split(' ', 1)
        first_name = parts[0]
        last_name = parts[1] if len(parts) > 1 else ''
        
        if len(t0.rows) > 1 and len(t0.rows[1].cells) > 5:
            t0.rows[1].cells[1].text = first_name
            t0.rows[1].cells[5].text = last_name or student.get('name', '')
        if len(t0.rows) > 2 and len(t0.rows[2].cells) > 3:
            t0.rows[2].cells[1].text = student.get('grade', "Β' Γυμνασίου")
            t0.rows[2].cells[3].text = student.get('class_section', "Β1")
        if len(t0.rows) > 3 and len(t0.rows[3].cells) > 6:
            t0.rows[3].cells[1].text = student.get('diagnosis', 'Ειδική Μαθησιακή Δυσκολία')
            t0.rows[3].cells[6].text = "2026-2027"

    os.makedirs(output_dir, exist_ok=True)
    st_clean_name = student.get('name', 'Μαθητής').replace(' ', '_')
    out_filename = f"4_Τελική_Έκθεση_{st_clean_name}.docx"
    out_path = os.path.join(output_dir, out_filename)
    doc.save(out_path)
    return out_path
